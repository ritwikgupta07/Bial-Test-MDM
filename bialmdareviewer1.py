import streamlit as st
import openai
from openai import AzureOpenAI
import json
import base64
import os
import pandas as pd
import tempfile
import html
import traceback
import re
import docx
import requests
from docx import Document
from docx.shared import Inches
import io
from dotenv import load_dotenv

# New imports for improved document download
import markdown
from bs4 import BeautifulSoup

# Azure Search imports
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery

# Langchain imports
from langchain_openai import AzureChatOpenAI as LangchainAzureChatOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain.schema import HumanMessage

load_dotenv()
# --- Page Config (Must be the first Streamlit command) ---
try:
    st.set_page_config(
        page_title="BIAL Multi-Agent Regulatory Platform", page_icon="✈️", layout="wide"
    )
except Exception as e_config:
    print(f"CRITICAL ERROR during st.set_page_config: {e_config}")
    st.error(f"Error during st.set_page_config: {e_config}")
    st.stop()


# --- Helper Function ---
def check_creds(cred_value, placeholder_prefix="YOUR_"):
    if not cred_value:
        return True
    if isinstance(cred_value, str):
        if (
            placeholder_prefix in cred_value.upper()
            or "ENTER_YOUR" in cred_value.upper()
            or (cred_value.startswith("<") and cred_value.endswith(">"))
        ):
            return True
    return False


# --- Main Application Logic ---
def main_app_logic():
    # --- Credentials Configuration ---
    AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
    AZURE_SEARCH_API_KEY = os.getenv("AZURE_SEARCH_API_KEY")
    DEFAULT_AZURE_SEARCH_INDEX_NAME = os.getenv("DEFAULT_AZURE_SEARCH_INDEX_NAME")
    DEFAULT_VECTOR_FIELD_NAME = os.getenv("DEFAULT_VECTOR_FIELD_NAME")
    DEFAULT_SEMANTIC_CONFIG_NAME = os.getenv("DEFAULT_SEMANTIC_CONFIG_NAME")

    AZURE_OPENAI_ENDPOINT_VAL = os.getenv("AZURE_OPENAI_ENDPOINT_VAL")
    AZURE_OPENAI_API_VERSION_VAL = os.getenv("AZURE_OPENAI_API_VERSION_VAL")
    AZURE_OPENAI_API_KEY_VAL = os.getenv("AZURE_OPENAI_API_KEY_VAL")
    DEPLOYMENT_ID_VAL = os.getenv("DEPLOYMENT_ID_VAL")
    PLANNING_LLM_DEPLOYMENT_ID = os.getenv("PLANNING_LLM_DEPLOYMENT_ID")
    SEARCH_QUERY_EMBEDDING_DEPLOYMENT_ID = os.getenv(
        "SEARCH_QUERY_EMBEDDING_DEPLOYMENT_ID"
    )

    BING_SEARCH_API_KEY = os.getenv("BING_SEARCH_API_KEY")
    BING_SEARCH_ENDPOINT = os.getenv("BING_SEARCH_ENDPOINT")
    SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY")

    # --- Initialize API Clients ---
    openai.api_type = "azure"
    openai.api_base = AZURE_OPENAI_ENDPOINT_VAL
    openai.api_version = AZURE_OPENAI_API_VERSION_VAL
    openai.api_key = AZURE_OPENAI_API_KEY_VAL

    search_query_embeddings_model, planning_openai_client, synthesis_openai_client = (
        None,
        None,
        None,
    )
    try:
        if not any(
            check_creds(c)
            for c in [
                AZURE_OPENAI_API_KEY_VAL,
                AZURE_OPENAI_ENDPOINT_VAL,
                SEARCH_QUERY_EMBEDDING_DEPLOYMENT_ID,
                AZURE_OPENAI_API_VERSION_VAL,
            ]
        ):
            search_query_embeddings_model = AzureOpenAIEmbeddings(
                azure_deployment=SEARCH_QUERY_EMBEDDING_DEPLOYMENT_ID,
                azure_endpoint=AZURE_OPENAI_ENDPOINT_VAL,
                api_key=AZURE_OPENAI_API_KEY_VAL,
                api_version=AZURE_OPENAI_API_VERSION_VAL,
                chunk_size=1,
            )
    except Exception as e:
        st.sidebar.error(f"Error initializing Embeddings Model: {e}")
    try:
        if not any(
            check_creds(c)
            for c in [
                AZURE_OPENAI_API_KEY_VAL,
                AZURE_OPENAI_ENDPOINT_VAL,
                AZURE_OPENAI_API_VERSION_VAL,
            ]
        ):
            if not check_creds(PLANNING_LLM_DEPLOYMENT_ID):
                planning_openai_client = AzureOpenAI(
                    api_key=AZURE_OPENAI_API_KEY_VAL,
                    azure_endpoint=AZURE_OPENAI_ENDPOINT_VAL,
                    api_version=AZURE_OPENAI_API_VERSION_VAL,
                )
            if not check_creds(DEPLOYMENT_ID_VAL):
                synthesis_openai_client = AzureOpenAI(
                    api_key=AZURE_OPENAI_API_KEY_VAL,
                    azure_endpoint=AZURE_OPENAI_ENDPOINT_VAL,
                    api_version=AZURE_OPENAI_API_VERSION_VAL,
                )
    except Exception as e:
        st.sidebar.error(f"Error initializing OpenAI Clients: {e}")

    # --- Tool Functions ---
    def query_bing_web_search(query: str, count: int = 5) -> str:
        if check_creds(BING_SEARCH_API_KEY) or check_creds(BING_SEARCH_ENDPOINT):
            return "Error: Bing Search API credentials are not configured."
        headers = {"Ocp-Apim-Subscription-Key": BING_SEARCH_API_KEY}
        params = {
            "q": query,
            "count": count,
            "textDecorations": True,
            "textFormat": "HTML",
        }
        try:
            response = requests.get(
                BING_SEARCH_ENDPOINT, headers=headers, params=params, timeout=10
            )
            response.raise_for_status()
            search_results = response.json()
            snippets = [
                f"Title: {res['name']}\nURL: {res['url']}\nSnippet: {res['snippet']}\n---"
                for res in search_results.get("webPages", {}).get("value", [])
            ]
            return (
                "\n".join(snippets) if snippets else "No Bing web search results found."
            )
        except Exception as e:
            return f"Error during Bing web search: {e}"

    def query_serpapi(query: str, count: int = 5) -> str:
        """Performs a web search using the SerpApi."""
        if check_creds(SERPAPI_API_KEY):
            return "Error: SerpApi API key is not configured."
        params = {
            "q": query,
            "api_key": SERPAPI_API_KEY,
            "num": count,
            "engine": "google",
        }
        try:
            response = requests.get(
                "https://serpapi.com/search.json", params=params, timeout=15
            )
            response.raise_for_status()
            search_results = response.json()
            organic_results = search_results.get("organic_results", [])
            snippets = [
                f"Title: {res.get('title', 'N/A')}\nURL: {res.get('link', 'N/A')}\nSnippet: {res.get('snippet', 'N/A')}\n---"
                for res in organic_results
            ]
            return (
                "\n".join(snippets)
                if snippets
                else "No web search results found via SerpApi."
            )
        except Exception as e:
            return f"Error during SerpApi search: {e}"

    def get_query_vector(text_to_embed):
        if not search_query_embeddings_model:
            st.toast("Search Query Embedding model not ready.", icon="⚠️")
            return None
        try:
            return search_query_embeddings_model.embed_query(text_to_embed)
        except Exception as e:
            st.error(f"Error generating query vector: {e}")
            return None

    @st.cache_data(ttl=3600)
    def get_indexes():
        indexes = (
            [DEFAULT_AZURE_SEARCH_INDEX_NAME]
            if DEFAULT_AZURE_SEARCH_INDEX_NAME
            and not check_creds(DEFAULT_AZURE_SEARCH_INDEX_NAME)
            else []
        )
        if check_creds(AZURE_SEARCH_API_KEY) or check_creds(AZURE_SEARCH_ENDPOINT):
            return list(set(indexes))
        try:
            search_creds = AzureKeyCredential(AZURE_SEARCH_API_KEY)
            index_client = SearchIndexClient(AZURE_SEARCH_ENDPOINT, search_creds)
            indexes.extend([index.name for index in index_client.list_indexes()])
        except Exception as e:
            st.sidebar.error(f"Cannot retrieve Azure Search indexes: {e}", icon="🚨")
        return list(set(indexes))

    def query_azure_search(
        query_text,
        index_name,
        k=5,
        use_hybrid_semantic_search=True,
        vector_field_name=DEFAULT_VECTOR_FIELD_NAME,
        semantic_config_name=DEFAULT_SEMANTIC_CONFIG_NAME,
    ):
        context, references_data = "", []
        if any(
            check_creds(c)
            for c in [AZURE_SEARCH_API_KEY, AZURE_SEARCH_ENDPOINT, index_name]
        ):
            return "Error: Azure Search credentials or index name are placeholders.", []
        try:
            search_client = SearchClient(
                AZURE_SEARCH_ENDPOINT,
                index_name,
                AzureKeyCredential(AZURE_SEARCH_API_KEY),
            )
            select_fields = ["content", "filepath", "url", "title"]
            search_kwargs = {
                "search_text": query_text if query_text and query_text.strip() else "*",
                "top": k,
                "include_total_count": True,
                "select": ",".join(select_fields),
            }
            if use_hybrid_semantic_search:
                if not check_creds(vector_field_name) and (
                    query_vector := get_query_vector(query_text)
                ):
                    search_kwargs["vector_queries"] = [
                        VectorizedQuery(
                            vector=query_vector,
                            k_nearest_neighbors=k,
                            fields=vector_field_name,
                        )
                    ]
                if not check_creds(semantic_config_name):
                    search_kwargs.update(
                        {
                            "query_type": "semantic",
                            "semantic_configuration_name": semantic_config_name,
                            "query_caption": "extractive",
                            "query_answer": "extractive",
                        }
                    )

            results = search_client.search(**search_kwargs)
            if results.get_count() == 0:
                return "", []

            processed_references = {}
            for idx, doc in enumerate(results):
                context += doc.get("content", "") + "\n\n"
                display_name = doc.get("title") or (
                    os.path.basename(doc.get("filepath", ""))
                    if doc.get("filepath")
                    else f"Source {idx+1}"
                )
                ref_key = doc.get("url") or display_name
                if ref_key not in processed_references:
                    processed_references[ref_key] = {
                        "filename_or_title": display_name,
                        "url": doc.get("url"),
                        "score": doc.get("@search.score"),
                        "reranker_score": doc.get("@search.reranker_score"),
                    }
            references_data = list(processed_references.values())
        except Exception as e:
            return f"Error accessing search index '{index_name}': {e}", []
        return context.strip(), references_data

    def get_query_plan_from_llm(user_question, client_for_planning):
        if not client_for_planning or check_creds(PLANNING_LLM_DEPLOYMENT_ID):
            return "Error: Planning LLM not configured.", None
        planning_prompt = f"""You are a query planning assistant specializing in breaking down complex questions about **AERA regulatory documents, often concerning tariff orders, consultation papers, control periods, and specific financial data (like CAPEX, Opex, Traffic) for airport operators such as DIAL, MIAL, BIAL, HIAL.**
Your primary task is to take a user's complex question related to these topics and break it down into a series of 1 to 10 simple, self-contained search queries that can be individually executed against a document index. Each search query should aim to find a specific piece of information (e.g., a specific figure, a justification, a comparison point) needed to answer the overall complex question.
If the user's question is already simple and can be answered with a single search, return just that single query in the list.
If the question is very complex and might require more distinct search steps, formulate the most critical 1 to 10 search queries, focusing on distinct pieces of information.
Return your response ONLY as a JSON list of strings, where each string is a search query.
User's complex question: {user_question}
Your JSON list of search queries:"""
        try:
            response = client_for_planning.chat.completions.create(
                model=PLANNING_LLM_DEPLOYMENT_ID,
                messages=[{"role": "user", "content": planning_prompt}],
                temperature=0.0,
                max_tokens=16000,
            )
            plan_str = response.choices[0].message.content
            if match := re.search(r"[.*]", plan_str, re.DOTALL):
                query_plan = json.loads(match.group(0))
                if isinstance(query_plan, list) and all(
                    isinstance(q, str) for q in query_plan
                ):
                    return None, query_plan
            return None, [user_question]
        except Exception as e:
            return f"Error getting query plan from LLM: {e}", None

    def generate_answer_from_search(
        user_question,
        index_name,
        use_hybrid_semantic,
        vector_field,
        semantic_config,
        temperature,
        max_tokens_param,
        client_for_synthesis,
        show_details=True,
        system_prompt_override=None,
    ):
        if not client_for_synthesis or check_creds(DEPLOYMENT_ID_VAL):
            return "Error: Synthesis LLM not configured."

        if show_details:
            st.write("⚙️ Generating query plan...")
        plan_error, query_plan = get_query_plan_from_llm(
            user_question, planning_openai_client
        )

        query_plan = query_plan or [user_question]
        if show_details:
            if plan_error:
                st.error(plan_error)
            st.write(f"📝 **Execution Plan:**")
            for i, q_step in enumerate(query_plan):
                st.write(f"   Step {i+1}: {q_step}")

        combined_context_for_llm = ""
        all_retrieved_details = []
        for i, sub_query in enumerate(query_plan):
            if show_details:
                st.write(f"🔍 Executing Step {i+1}: Searching for '{sub_query}'...")
            context_for_step, retrieved_details_for_step = query_azure_search(
                sub_query,
                index_name,
                use_hybrid_semantic_search=use_hybrid_semantic,
                vector_field_name=vector_field,
                semantic_config_name=semantic_config,
            )
            if context_for_step and not context_for_step.startswith("Error"):
                combined_context_for_llm += (
                    f"\n\n--- Context for sub-query: '{sub_query}' ---\n"
                    + context_for_step
                )
                all_retrieved_details.extend(retrieved_details_for_step)

        if show_details:
            st.session_state.last_retrieved_chunks_search = all_retrieved_details

        if not combined_context_for_llm.strip():
            return "No relevant information found in the search index across all planned queries."

        formatted_sources = []
        if all_retrieved_details:
            unique_sources_dict = {
                (item.get("url") or item.get("filename_or_title")): item
                for item in all_retrieved_details
            }
            for i, item in enumerate(list(unique_sources_dict.values())):
                source_line = f"Source {i+1}: {html.escape(item['filename_or_title'])}"
                if item.get("url"):
                    source_line += f" ([Link]({html.escape(item['url'])}))"
                if item.get("score") is not None:
                    source_line += f" [Score: {item['score']:.4f}]"
                if item.get("reranker_score") is not None:
                    source_line += f" [Reranker Score: {item['reranker_score']:.4f}]"
                formatted_sources.append(source_line)
        else:
            formatted_sources.append(
                "No specific source metadata retrieved, but context was found."
            )
        formatted_refs_str = "\n".join(formatted_sources)

        if show_details:
            st.write("💡 Synthesizing final answer...")

        user_base_prompt_instructions = (
            system_prompt_override
            or """Part 1:   "you are AI assistant for Multiyear tarrif submission for AERA. final response should have 1500 words at least.
* **3.2. Source Attribution (Authority Data - Handling Terminology for Authority's Stance):**
    If the query relates to the "Authority's" stance (e.g., user asks “What is the authority’s *approved* change?”, “What did the authority *decide*?”, “What was *proposed* by the authority?”), your primary goal is to find the authority's documented action or position on that specific subject *within the  .
    * **Understanding Levels of Finality (for AI's internal logic):**
        * **Final/Conclusive Terms:** "approved by the authority", "decided by the authority", "authority's decision", "final tariff/order states", "sanctioned by authority", "adopted by authority".
        * **Provisional/Draft Terms:** "proposed by theauthority", "authority's proposal", "draft figures", "recommended by the authority" (if it's a recommendation for a later decision).
        * **Analytical/Consideration Terms:** "considered by the authority", "analyzed by the authority", "examined by the authority", "authority's review/view/assessment/preliminary findings".
    * **Extraction Strategy Based on User Intent and Document Content:**
        1.  **Attempt to Match User's Exact Term First:** Always search the CONTEXT for information explicitly matching the user's specific terminology (e.g., if the user asks for "approved," look first for "approved by the authority"). If found, present this.
        2.  **If User's Query Implies Finality (e.g., asks for "approved," "final figures," "decision"):**
            * And their *exact term* is NOT found in the CONTEXT for that item:
                * **Prioritize searching the CONTEXT for other Final/Conclusive terms** (e.g., "decided by the authority," "authority's decision"). If one of these is found, present this. You MUST then state clearly: "You asked for 'approved.' The document describes what was '*[actual term found, e.g., decided by the authority]*' as follows: [data and references]."
                * If no Final/Conclusive terms are found for that item in the CONTEXT, then (and only then) look for Provisional/Draft terms (e.g., "proposed by the authority"). If found, present this, stating: "You asked for 'approved.' A final approval or decision was not found for this item in the provided context. However, the authority '*proposed*' the following: [data and references]."
                * If neither of the above is found, look for Analytical/Consideration terms and report similarly with clarification.
        3.  **If User's Query Uses a Provisional/Draft Term (e.g., "proposed"):**
            * Prioritize finding information matching those Provisional/Draft terms in the CONTEXT.
            * If not found, you can then look for Analytical/Consideration terms, clarifying the terminology. Avoid presenting Final/Conclusive terms unless you explicitly state that the user asked for a draft but a final version was found regarding that specific point.
        4.  **If User's Query Uses an Analytical/Consideration Term (e.g., "considered"):**
            * Prioritize finding information matching those terms. Always give me the table refrence for the response i mean which table no you have refered for the response
    * **Accurate Reporting is Key:** Always present information using the **document's actual terminology**. Clearly explain if and how it relates to the user's original query terms. If multiple relevant stages of authority action are evident in the CONTEXT (e.g., a proposal and then a later decision), you may summarize both, clearly distinguishing them by the terms used in the document.
    * **If No Relevant Authority Action Found:** If the provided CONTEXT contains no clear information matching any relevant stage of authority action (final, provisional, or analytical) regarding the specific subject of the query, state that this information was not found for the authority in the provided context.
    * **Table Headers:** Use data from tables if their headers clearly indicate the source and nature of the data (e.g., "Figures as Decided by the Authority," "Operator's Proposed Traffic").
"**Crucial Instruction for Authority's Stance:** When the user asks for 'approved', 'final', or 'decided' figures from the Authority, it is **imperative** that you prioritize extracting text explicitly labeled with terms like 'decided by the authority' or 'approved by the authority' from the CONTEXT. If such conclusive terms are present for the queried item, present them. Only if NO such conclusive terms are found in the CONTEXT for that item should you then present information labeled 'proposed by the authority', and you MUST clearly state that you are providing 'proposed' figures because 'approved/decided' ones were not found in the given context."
* **Clarifying "Considered by Authority" in a "True-Up" Context:**
    * If the user query asks what the authority 'considered' in relation to a 'true-up' of a specific control period:
        * First, check if the CONTEXT contains information about the authority's **analysis or verification of the actual figures submitted for that true-up period**. If so, present this.
        * If the user's query might also imply understanding the **original baseline** that is being trued-up against, you can additionally (or if the above is not found) look for what the authority **originally considered or determined when setting the tariffs for that control period at the beginning of that period**.
        * **Crucially, always differentiate these two.** For example: "For the true-up of the Third Control Period, DIAL submitted the following actual traffic figures (e.g., from Table 25): [data]. The figures that the Authority had originally considered at the time of determining the tariff for the Third Control Period were (e.g., from Table 26): [data]."
        * If the query is simply "What was considered for true-up..." without specifying "original determination" vs "actuals review", and both types of information are in the context, you might offer both or ask the user to clarify which aspect of "considered for true-up" they are interested in Mandatory Table Referencing:
For any data, figures, or claims extracted from a table within the CONTEXT, you must cite the corresponding table number in your response. This is a strict requirement for all outputs. The reference should be placed directly with the data it pertains to.
Example 1: "The Authority approved Aeronautical Revenue of ₹1,500 Cr for FY 2024-25 (Table 15)."
Example 2: "For the true-up, the operator submitted actual passenger traffic of 45 million (as per Table 3.2), while the original figure considered by the Authority was 42 million (from Table 5.1 of the original Order)."""
        )

        synthesis_prompt = (
            f"{user_base_prompt_instructions}\n\n"
            f"Based on the general background above, please synthesize a comprehensive answer to the original USER QUESTION using all the following retrieved CONTEXT from multiple search steps and the IDENTIFIED CONTEXT SOURCES.\n\n"
            f"ORIGINAL USER QUESTION: {user_question}\n\n"
            f"AGGREGATED CONTEXT (from multiple search steps):\n---------------------\n{combined_context_for_llm}\n---------------------\n\n"
            f"IDENTIFIED CONTEXT SOURCES (from metadata):\n---------------------\n{formatted_refs_str}\n---------------------\n\n"
            f"SPECIFIC INSTRUCTIONS FOR YOUR RESPONSE (in addition to the general background provided):\n"
            f"1. Directly address all parts of the ORIGINAL USER QUESTION.\n"
            f"2. Synthesize information from the different context sections if they relate to different aspects of the original question.\n"
            f"3. Format numerical data extracted from tables into an HTML table with borders (e.g., <table border='1'>...). Use table headers (<th>) and table data cells (<td>).\n"
            f"4. **References are crucial.** At the end of your answer, include a 'References:' section listing the source documents (using filenames or titles as provided in 'IDENTIFIED CONTEXT SOURCES') from which the information was derived. If a URL is available for a source, make the filename/title a clickable hyperlink to that URL.\n\n"
            f"COMPREHENSIVE ANSWER TO THE ORIGINAL USER QUESTION:\n"
            f"*Crucially:* Include references for the information presented. Mention the specific source (e.g., the filename from the IDENTIFIED CONTEXT SOURCES list) and, if mentioned within the text context itself, include table numbers (e.g., 'Table 26') or section titles and the file name. Present these references clearly at the end of your answer under a 'References:' heading."
            "\n\nANSWER:"
        )
        try:
            response = client_for_synthesis.chat.completions.create(
                model=DEPLOYMENT_ID_VAL,
                messages=[{"role": "user", "content": synthesis_prompt}],
                temperature=temperature,
                max_tokens=max_tokens_param,
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating synthesized answer: {e}"

    def refine_and_regenerate_report(
        original_report: str, new_info: str, client_for_synthesis
    ) -> str:
        if not client_for_synthesis:
            return "Error: Synthesis LLM client not initialized."
        refinement_prompt = f"""You are a report writing expert. Your task is to seamlessly integrate a new piece of information into an existing report. Do not simply append the new information. Instead, find the most relevant section in the 'ORIGINAL REPORT' and intelligently merge the 'NEW INFORMATION' into it. Rewrite paragraphs as needed to ensure the final report is coherent, clean, and well-integrated. Return ONLY the full, updated report text.
        **ORIGINAL REPORT:**
        ---
        {original_report}
        ---
        **NEW INFORMATION TO INTEGRATE:**
        ---
        {new_info}
        ---
        **FULL, REFINED, AND INTEGRATED REPORT:**"""
        with st.spinner("✨ Refining report with new information..."):
            try:
                response = client_for_synthesis.chat.completions.create(
                    model=DEPLOYMENT_ID_VAL,
                    messages=[{"role": "user", "content": refinement_prompt}],
                    temperature=0.2,
                    max_tokens=st.session_state.conv_agent_max_tokens,
                )
                st.toast("Report successfully refined!", icon="✅")
                return response.choices[0].message.content
            except Exception as e:
                st.error(f"Error during report refinement: {e}")
                return original_report

    def extract_text_from_docx(file):
        try:
            return "\n".join([para.text for para in docx.Document(file).paragraphs])
        except Exception as e:
            st.error(f"Error reading Word document: {e}")
            return None

    def parse_html_to_docx(soup, document):
        """Recursively parses BeautifulSoup elements and adds them to a docx document."""
        for element in soup.find_all(True, recursive=False):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                try:
                    level = int(element.name[1])
                    document.add_heading(element.get_text(strip=True), level=level)
                except (ValueError, IndexError):
                    document.add_heading(element.get_text(strip=True), level=2)
            elif element.name == "p":
                document.add_paragraph(element.get_text(strip=True))
            elif element.name == "table":
                try:
                    rows = element.find_all("tr")
                    if not rows:
                        continue

                    headers = [
                        th.get_text(strip=True) for th in rows[0].find_all(["th", "td"])
                    ]
                    if not headers:
                        continue

                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header

                    for row in rows[1:]:
                        row_cells_data = [
                            td.get_text(strip=True) for td in row.find_all("td")
                        ]
                        if len(row_cells_data) == len(headers):
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(row_cells_data):
                                row_cells[i].text = cell_text
                except Exception as e:
                    document.add_paragraph(f"(Error parsing table: {e})")
            elif element.name in ["ul", "ol"]:
                style = "List Bullet" if element.name == "ul" else "List Number"
                for li in element.find_all("li", recursive=False):
                    p = document.add_paragraph(style=style)
                    p.add_run(li.get_text(strip=True, separator=" ").split("\n")[0])
            elif element.name in ["div", "section", "article"]:
                parse_html_to_docx(element, document)

    def create_word_document(markdown_text):
        """
        Creates a Word document from markdown text, correctly parsing HTML tags.
        """
        try:
            document = Document()
            styles = document.styles
            try:
                if "List Bullet" not in styles:
                    styles.add_style("List Bullet", 1).base_style = styles[
                        "List Paragraph"
                    ]
                if "List Number" not in styles:
                    styles.add_style("List Number", 1).base_style = styles[
                        "List Paragraph"
                    ]
            except Exception as e:
                print(f"Could not add default list styles: {e}")

            html_content = markdown.markdown(
                markdown_text, extensions=["markdown.extensions.tables"]
            )
            soup = BeautifulSoup(html_content, "html.parser")

            parse_html_to_docx(soup, document)

            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            return file_stream
        except Exception as e:
            st.error(f"Failed to create Word document: {e}")
            document = Document()
            document.add_paragraph("Error creating document.")
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            return file_stream

    @st.cache_data
    def get_base64_image(image_path: str):
        try:
            with open(image_path, "rb") as f:
                return base64.b64encode(f.read()).decode("utf-8")
        except Exception:
            return None

    # --- UI & State ---
    st.markdown(
        """ <style>
    :root {
        --primary-background: #33263C; /* Dominant dark purple/brown background */
        --card-background: #2A2A2A; /* Dark gray of the card/table header */
        --text-color-light: #E0E0E0; /* Light gray for most text */
        --text-color-dark: #A0A0A0; /* Slightly darker gray for some text or secondary information */
        --header-text-color: #FFFFFF; /* White for the header "Deploy" and column titles */
        --button-green: #28A745; /* Green color of the "Generate" button */
        --sidebar-dark: #1E1E1E; /* Very dark almost black for the sidebar area */
    }

    /* General Styles */
    body {
        color: var(--text-color-light);
    }
    .stApp {
        background-color: var(--primary-background);
    }

    /* Main content area (response box, text area, etc.) */
    .response-box, .stTextArea, .stTextInput, .st-expander {
        background-color: var(--card-background) !important;
        border: 1px solid #4A4A4A !important; /* A slightly lighter border for definition */
        color: var(--text-color-light) !important;
        border-radius: 8px;
    }
    .st-expander header {
        color: var(--text-color-light) !important;
    }

    .response-box {
        padding: 20px;
        margin-bottom: 20px;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    
    /* History Entry in Sidebar */
    .history-entry {
        border: 1px solid #4A4A4A;
        padding: 15px;
        margin-bottom: 15px;
        border-radius: 8px;
        background-color: var(--card-background);
        white-space: pre-wrap;
        word-wrap: break-word;
    }

    /* Titles and Headers */
    .title-bar h1, h1, h2, h3, h4, p {
        color: var(--header-text-color);
    }
    
    /* --- CSS FOR ALL BUTTONS TO BE GREEN --- */
    .stButton>button {
        border: none;
        background-color: var(--button-green);
        color: var(--header-text-color);
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: bold;
        transition: all 0.3s ease-in-out;
    }

    .stButton>button:hover {
        background-color: #218838; /* A darker green for hover */
        box-shadow: 0 0 15px rgba(40, 167, 69, 0.5);
    }
    
    .stButton>button:active {
        background-color: #1E7E34 !important; /* An even darker green for active/click */
    }

    div[data-testid="column"] .stButton>button {
        height: 100%;
        width: 100%;
        text-align: left !important;
        padding: 15px;
        font-weight: normal;
    }

    div[data-testid="column"] .stButton>button:hover {
        transform: translateY(-3px);
    }
    
    /* --- END OF BUTTON CSS --- */

    .st-emotion-cache-1r6slb0 { /* Sidebar */
        background-color: var(--sidebar-dark);
    }
    .st-emotion-cache-16txtl3 { /* Main content area */
        background-color: var(--primary-background);
    }
    
    </style>
    
    
    """,
        unsafe_allow_html=True,
    )

    default_session_state = {
        "conversation_history": [],
        "last_retrieved_chunks_search": [],
        "question_text": "",
        "conv_agent_temp": 0.5,
        "conv_agent_max_tokens": 16000,
        "conv_agent_selected_index": DEFAULT_AZURE_SEARCH_INDEX_NAME,
        "conv_agent_use_hybrid": True,
        "conv_agent_vector_field": DEFAULT_VECTOR_FIELD_NAME,
        "conv_agent_semantic_config": DEFAULT_SEMANTIC_CONFIG_NAME,
        "app_mode": "Conversational Agent",
        "mda_analysis_type": "MDA Manpower Analysis",
        "mda_report_content": None,
        "mda_chat_history": [],
        "web_search_engine": "Bing",  # Default search engine
    }
    for key, value in default_session_state.items():
        if key not in st.session_state:
            st.session_state[key] = value

    with st.sidebar:
        st.sidebar.title("Choose App")
        st.session_state.app_mode = st.sidebar.radio(
            "Select the application to use:",
            ("Conversational Agent", "MDA Reviewer"),
            key="app_mode_selector",
        )
        st.markdown("---")
        with st.expander("⚙️ Settings & History", expanded=True):
            st.header("Agent Settings")
            st.session_state.conv_agent_temp = st.slider(
                "Temperature", 0.0, 1.0, st.session_state.conv_agent_temp, 0.1
            )
            st.session_state.conv_agent_max_tokens = st.slider(
                "Output", 100, 16000, st.session_state.conv_agent_max_tokens, 50
            )

            st.session_state.web_search_engine = st.radio(
                "Web Search Engine", ("Bing", "Google (via Serper)")
            )

            st.subheader("Regulatory database")
            available_indexes = get_indexes()
            st.session_state.conv_agent_selected_index = st.selectbox(
                "Select Index",
                available_indexes,
                index=(
                    0
                    if not st.session_state.conv_agent_selected_index
                    in available_indexes
                    else available_indexes.index(
                        st.session_state.conv_agent_selected_index
                    )
                ),
            )
            st.session_state.conv_agent_use_hybrid = st.checkbox(
                "Enable Multi-Agent Flow", value=st.session_state.conv_agent_use_hybrid
            )
            if st.session_state.conv_agent_use_hybrid:
                st.session_state.conv_agent_vector_field = st.text_input(
                    "Vector Field Name", value=st.session_state.conv_agent_vector_field
                )
                st.session_state.conv_agent_semantic_config = st.text_input(
                    "Semantic Configuration",
                    value=st.session_state.conv_agent_semantic_config,
                )

            if st.session_state.app_mode == "Conversational Agent":
                st.markdown("---")
                if st.session_state.conversation_history:
                    st.subheader("📜 Chat History")
                    for entry in reversed(st.session_state.conversation_history):
                        st.markdown(
                            f"**Q:** {entry['question']}", unsafe_allow_html=True
                        )
                        st.markdown(
                            f"<div class='history-entry'>{entry['answer']}</div>",
                            unsafe_allow_html=True,
                        )
                        st.markdown("---", unsafe_allow_html=True)
                    if st.button("Clear History"):
                        st.session_state.conversation_history = []
                        st.session_state.question_text = ""
                        st.rerun()

    col_header1, col_header2 = st.columns([1, 10])
    with col_header1:
        if main_logo_base64 := get_base64_image("bial_logo.png"):
            st.image(f"data:image/png;base64,{main_logo_base64}", width=400)
    with col_header2:
        st.markdown(
            '<div class="title-bar"><h1>BIAL Regulatory Assistant</h1></div>',
            unsafe_allow_html=True,
        )
    st.markdown(
        "<hr style='margin-top: 0; margin-bottom:1em;'>", unsafe_allow_html=True
    )

    if st.session_state.app_mode == "Conversational Agent":
        st.subheader(
            f"💬 Conversational Agent (Index: {st.session_state.conv_agent_selected_index or 'Not Selected'})"
        )
        if st.session_state.last_retrieved_chunks_search:
            with st.expander("View Retrieved Context Sources", expanded=False):
                for item in st.session_state.last_retrieved_chunks_search:
                    st.markdown(
                        f"**{html.escape(item.get('filename_or_title', 'Source'))}** ..."
                    )

        st.markdown("---")
        st.subheader("Ask a question")
        predefined_questions = [
            "Calculate and compare the YoY change of employee expenses of DIAL and MIAL for the fourth control period",
            "What is the YoY change of employee expenses submitted by MIAL for the fourth control period and the rationale for the growth rates",
            "Compare the manpower expense per total passenger traffic submitted by DIAL and MIAL respectively for fourth control period.",
        ]
        cols = st.columns(len(predefined_questions))
        for i, q in enumerate(predefined_questions):
            if cols[i].button(q, key=f"predef_q_{i}"):
                st.session_state.question_text = q

        def update_question_text_conv():
            st.session_state.question_text = (
                st.session_state.question_text_area_main_common
            )

        st.text_area(
            "Your question:",
            value=st.session_state.question_text,
            key="question_text_area_main_common",
            on_change=update_question_text_conv,
        )

        submit_button_main = st.button("Submit Question", type="primary")

        if submit_button_main:
            if user_query := st.session_state.question_text.strip():
                answer = None
                with st.spinner("Thinking..."):
                    answer = generate_answer_from_search(
                        user_question=user_query,
                        index_name=st.session_state.conv_agent_selected_index,
                        use_hybrid_semantic=st.session_state.conv_agent_use_hybrid,
                        vector_field=st.session_state.conv_agent_vector_field,
                        semantic_config=st.session_state.conv_agent_semantic_config,
                        temperature=st.session_state.conv_agent_temp,
                        max_tokens_param=st.session_state.conv_agent_max_tokens,
                        client_for_synthesis=synthesis_openai_client,
                        show_details=True,
                    )
                if answer:
                    st.session_state.conversation_history.append(
                        {"question": user_query, "answer": answer}
                    )
                    st.subheader("💡 Answer")
                    st.markdown(
                        f'<div class="response-box">{answer}</div>',
                        unsafe_allow_html=True,
                    )
            else:
                st.warning("Please enter a question.")

    elif st.session_state.app_mode == "MDA Reviewer":
        st.subheader("📄 Review and validation of MDA")
        st.info(
            "Analyze and validate BIAL's regulatory submissions by benchmarking against peer airports and historical data.",
            icon="ℹ️",
        )

        analysis_prompts_config = {
            "MDA Manpower Analysis": {
                "Analysis of manpower expenditure projection for BIAL for fourth control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the manpower expenditure projected by BIAL: 1- Year on Year growth of personnel cost projected by BIAL for fourth control period in personnel cost_MDA and projections.pdf file. 2- Justification for personnel cost growth in fourth control period provided by BIAL as perpersonnel cost_MDA and projections.pdf file . 3- year on year Manpower Expenses growth Submitted by DIAL for fourth control period in DIAL fourth control period consultation Paper. 4- Justification provided by DIAL for manpower expenses submitted by DIAL for fourth control period. 5- Examination and rationale provided by authority for manpower expenses submitted by DIAL for fourth control period. 6- Year on Year growth of employee cost submitted by MIAL for fourth control period for fourth control period in MIAL Fourth control consultation Paper. 7- Justification provided by MIAL for manpower expenses per passeneger traffic submitted by MIAL for fourth control period. 8- Examination and rationale provided by authority for manpower expenses submitted by MIAL for fourth control period. 9- Using the rationale extracted in steps 4, 5 7 and 8 suggest how the rationale or justification provided by BIAL in the MDA document for manpower expenditure for fourth control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL tariff orders or consultation papers.",
                "Analysis of actual manpower expenditure for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the actual manpower expenditure for the third control period: 1. Actual manpower expenditure for BIAL and variance from authority approved manpower expenditure for the third control period. 2. Justification for manpower expenditure in third control period provided by BIAL. 3 Actual manpower expenditure for DIAL and variance from authority approved manpower expenditure for the third control period. 4. Justification provided by DIAL for actual manpower expenses for third control period and the reason for variance compared to authority approved figures. 5. Examination and rationale provided by authority for actual manpower expenditure for only DIAL for third control period and its variance compared to authority approved figures. 6. Actual manpower expenditure for MIAL and variance from authority approved manpower expenditure for the third control period. 7. Justification provided by MIAL for actual manpower expenses submitted by MIAL for third control period and the reason for variance with authority approved figures. 8. Examination and rationale provided by authority for actual manpower expenditure submitted by only MIAL for third control period and its variance compared to authority approved figures. 9. Using the rationale extracted in steps 4, 5, 7, and 8, suggest how the rationale or justification provided by BIAL in the MDA document for manpower expenditure for the third control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL tariff orders or consultation papers.",
                "Analysis of KPI Computation for BIAL for fourth Control period": f"the upload document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the KPI Computation.Calculate and compare the YoY change of employee expenses of DIAL and MIAL for the fourth control period,first give what is total manpower expense submitted by DIAL for fourth control period , employee cost submitted by MIAL for fourth control period . after wards calculate the passanger traffic submitted by DIAL and MIAL for fourth control period . divide the passenger traffic per manpoer cost anf compare it anf give us the rationale . Step 1: KPI Comparison. To begin, you will collect specific data from the DIAL Fourth Control Period Consultation Paper and DIAL Fourth Control Period Tariff Order, as well as the MIAL Fourth Control Period Consultation Paper and MIAL Fourth Control Period Tariff Order. From these documents, meticulously extract the manpower count, total passenger traffic, and total manpower expenditure for each fiscal year of their respective fourth control periods. With this comprehensive dataset, proceed to calculate two critical KPIs for both airports: manpower count per total passenger traffic and manpower expenditure per total passenger traffic. Once these KPIs are computed, compare them to BIAL's corresponding figures, assessing whether BIAL’s KPIs are higher, lower, or in line, while being careful to only compare data for years where the passenger traffic is similar to ensure the KPI comparison is accurate and meaningful. First, carefully examine BIAL's provided MDA document to identify the specific justifications for its manpower expense projections, including any explanations for variances from the prior control period. Next, to enhance this rationale, you will consult the detailed analyses and findings in the DIAL and MIAL Fourth Control Period Consultation Papers and Tariff Orders. Specifically, you will look for how these regulatory documents justify their own employee expense projections, such as by detailing factors like inflation, annual growth rates, and specific manpower growth factors tied to strategic operational expansions. Using these as a benchmark, you will then suggest improvements for BIAL's own justifications, for example, by recommending that BIAL provide a more granular breakdown of cost drivers, link employee growth to new projects or terminal expansions, or justify its average cost per employee based on specific salary benchmarks or industry-wide trends, ultimately making BIAL's rationale as transparent and well-supported as OF DIAL and MIAL.",
                "comparison between DIAL,MIAL and BIAL On Manpower expenses for third and fourth control period": f"""The uploaded document proposes the following: '{{document_summary}}'. Conduct a detailed comparative analysis of manpower expenses for BIAL, DIAL, and MIAL across the third and fourth control periods. **1.table presents a comparison of projected Year-on-Year (YoY) growth rates for total manpower expenditure, total headcount, and resulting cost per employee for BIAL, DIAL, and MIAL during the fourth control period . * **Justification Analysis:** Juxtapose the key drivers and rationale provided by each airport for their projected growth. Identify the strengths and weaknesses in the evidence **BIAL** has provided compared to its peers. * **KPI Benchmarking:**need a comparison table for employee cost per passenger traffic and  manpower count per passenger traffic submitted by BIAL ,DIAL and MIAL for fourth control period   *actuals* of **DIAL** and **MIAL** from years they handled similar traffic volumes. **2. Third Control Period (Actuals Analysis)**  need the comparative  table for The variance between authority-approved and actual figures for total expenditure and headcount across BIAL, DIAL, and MIAL during the third control period. * **Rationale Comparison:** Analyze the reasons provided by each airport for their variances. Critically, note which of the justifications from **DIAL** and **MIAL** were ultimately accepted or rejected by the regulatory authority in the final true-up orders. **3. Synthesis and Recommendations for BIAL** Synthesize the findings to formulate specific recommendations for **BIAL**: * **Identify Divergences:** Pinpoint every key metric where **BIAL's** projections or past performance significantly diverges from the benchmarks set by **DIAL** and **MIAL**. * **Formulate Enhancements:** For each divergence, suggest a concrete enhancement to **BIAL's** rationale, directly referencing the stronger justification, superior performance, or accepted regulatory precedent demonstrated by **DIAL** or **MIAL**. """,
            },
            "Utility Analysis": {
                "Analysis of electricity expenditure projection for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the electricity expenditure projected by BIAL: 1- Year on Year actual growth of power consumption cost for third control period  by BIAL in Utlities_cost_MDA document . 2-   Year on Year  actual power consumption by BIAL submitted in the Utlities_cost_MDA document. 3- Year on Year actual recoveries of power consumption by BIAL for third control period  in the Utlities_cost_MDA document. 4- Justification provided by BIAL for the power expense  and the variance of power expense with authority approved figures in third control period in the Utlities cost_MDA document. 5- Year on Year growth of actual power expense submitted by DIAL for true up of third control period in the fourth control period consultation paper. 6- Year on Year  growth of power consumption submitted by DIAL for third control period in the fourth control period consultation paper. 7- Year on Year actual recoveries from sub-concessionaries (%) submitted by DIAL for third control period in the fourth control period consultation paper. 8- Justification for actual power expense in third control period provided by DIAL and the variance with authority approved figures in fourth control period consultation paper. 9- Examination and rationale provided by authority on actual power cost and consumption submitted by DIAL for third control period in the fourth control period consultation paper.  10- Year on Year  Electricity cost(utility expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 11- Year on Year  electricity  gross consumption(utlity expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 12- Year on Year  recoveries of electricity consumption submitted by MIAL for the trueup of third control period in the MIAL fourth control period consultation paper. 8 Justification for actual electricity cost for the true up of third control period provided by MIAL in the MIALfourth control period consultation paper and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual Electricity cost and consumption submitted by MIAL true of third control period in the MIAL fourth control period consultation paper.15- Using the rationale extracted in steps 4, 8, 9,13 and 14 suggests how the rationale or justification provided by BIAL in the MDA document for electricity cost  for third control period can be enhanced. For every suggestion made, give specific reason why the suggestion was made using relevant references from DIAL and MIAL tariff orders or consultation papers. when asked about MIAL only give information relevant to MIAL not DIAL Strictly.",
                "Analysis of water expenditure projection for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the water expenditure projected by BIAL: 1-  actual portable and raw water cost by BIAL for trueup of third control period in Utlities_cost_MDA document . 2-year on Year raw and portable water  consumption by BIAL of true up for third control period in the Utilities cost_MDA document . 3- Year on Year actual recoveries of  water consumption by BIAL for the third control period in the Utlities_cost_MDA document . 4- Justification provided by BIAL for the water cost for third control period and the variance of water expense with authority approved figures in third control period in the Utlities_cost_MDA document. 5- Year on Year  water gross charge submitted by DIAL for true up of third control period in the DIAL fourth control period consultation paper. 6- Year on Year growth of water consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 7- Year on Year actual recoveries from sub- concessionaire submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 8- Justification for actual  gross water charge  in third control period in the DIAL fourth control period consultation paper provided by DIAL and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual water gross charge and consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper.  10- Year on Year water expense(utility expenses) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 11- Year on Year water consumption(Kl) submitted by MIAL for true up of third control period in the MIAL fourth control period consultation paper. 12- Year on Year  recoveries(kl) of water consumption submitted by MIAL for true up of the  third control period in the MIAL fourth control period consultation paper. 8- Justification for actual water gross amount for third control period in the MIAL fourth control period consultation paper provided by MIAL and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual water gross amount  and consumption submitted by MIAL for third control period in the MIAL fourth control period consultation paper.15- Using the rationale extracted in steps 4, 8, 9,13 and 14 suggest how the rationale or justification provided by BIAL in the MDA document for water expenditure for trueup of  third control period can be enhanced. For every suggestion made, give specific reason why the suggestion made using relevant references from DIAL and MIAL tariff orders or consultation papers.",
                "Analysis of KPI Computation for BIAL(Utility Expenditure)": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analyzing the KPI Computation. Calculate and compare the YoY change of power and electricity e expenses of DIAL and MIAL for true up of third control period p,first give what is total electricity expense submitted by DIAL for true up of third control period  , Electricty cost submitted by MIAL for true up of third control  period . after wards calculate the passanger traffic submitted by DIAL and MIAL for true up of third control period divide the passenger traffic per electricity cost  and compare it and give us the rationale ,Year on Year  water gross charge submitted by DIAL per passenger traffic submitted  for true up of third control period in the DIAL fourth control period consultation paper. Calculate and compare the YoY change of water a gross charge of DIAL and MIAL for true up of third control period p,first give what is total electricity expense submitted by DIAL for true up of third control period  , water cost submitted by MIAL for true up of third control  period . after wards calculate the passanger traffic submitted by DIAL and MIAL for true up of thord control perioddivide the passenger traffic per water cost and compare it and give us the rationale Step 1: KPI Comparison. To begin, you will collect specific data from the DIAL Fourth Control Period Consultation Paper and DIAL Fourth Control Period Tariff Order, as well as the MIAL Fourth Control Period Consultation Paper and MIAL Fourth Control Period Tariff Order. From these documents, meticulously extract the electricity consumption, water consumption, and total passenger traffic for each fiscal year of their respective fourth control periods. With this comprehensive dataset, proceed to calculate two critical KPIs for both airports: electricity consumption per total passenger traffic and water consumption per total passenger traffic. Once these KPIs are computed, compare them to BIAL's corresponding figures, assessing whether BIAL's KPIs are higher, lower, or in line, while being careful to only compare data for years where the passenger traffic is similar to ensure the KPI comparison is accurate and meaningful. First, carefully examine BIAL's provided MDA document to identify the specific justifications for its utility expense projections, including any explanations for variances from the prior control period. Next, to enhance this rationale, you will consult the detailed analyses and findings in the DIAL and MIAL Fourth Control Period Consultation Papers and Tariff Orders. Specifically, you will look for how these regulatory documents justify their own utility expense projections, such as by detailing factors like energy efficiency initiatives, water conservation projects, infrastructure upgrades impacting consumption, or changes in operational scope. Using these as a benchmark, you will then suggest improvements for BIAL's own justifications, for example, by recommending that BIAL provide a more granular breakdown of consumption drivers, link utility usage to new terminal operations or technological advancements, or justify its per-passenger consumption figures based on industry best practices or environmental targets, ultimately making BIAL rationale as transparent and well supported as that of its peers.",
                "comparison between DIAL,MIAL and BIAL on utility  expenses for third control period": f"""The uploaded document proposes the following: '{{document_summary}}'. Conduct a detailed comparative analysis of utility expenses for BIAL, DIAL, and MIAL across the third and fourth control periods. **table presents a comparison of projected Year-on-Year (YoY) growth rates for total Electricity cost, Actual electricity consumption, and  for Year on Year actual recoveries from sub-concessionaries (%) for BIAL, DIAL, and MIAL during the third  control period . * **Justification Analysis:** Juxtapose the key drivers and rationale provided by each airport for their projected growth. Identify the strengths and weaknesses in the evidence **BIAL** has provided compared to its peers. * **KPI Benchmarking:**need a comparison table for Electricity Consumption per total passenger traffic  and Water consumption per total passenger traffic  by BIAL ,DIAL and MIAL for third control period   *actuals* of **DIAL** and **MIAL** from years they handled similar traffic volumes. **2. Third Control Period (Water)**  need the comparative  table for The variance between authority-approved and actual figures for total water cost and actual water across BIAL, DIAL, and MIAL during the third control period. * **Rationale Comparison:** Analyze the reasons provided by each airport for their variances. Critically, note which of the justifications from **DIAL** and **MIAL** were ultimately accepted or rejected by the regulatory authority in the final true-up orders. **3. Synthesis and Recommendations for BIAL** Synthesize the findings to formulate specific recommendations for **BIAL**: * **Identify Divergences:** Pinpoint every key metric where **BIAL's** projections or past performance significantly diverges from the benchmarks set by **DIAL** and **MIAL**. * **Formulate Enhancements:** For each divergence, suggest a concrete enhancement to **BIAL's** rationale, directly referencing the stronger justification, superior performance, or accepted regulatory precedent demonstrated by **DIAL** or **MIAL**. """,
            },
            "R&M Analysis": {
                "Analysis of repairs and maintenance expenditure for true up for BIAL for third control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the repairs and maintenance expenditure projected by BIAL: 1- Year on Year actual growth of repairs and maintenance expenditure by BIAL for third control period in the MDA_O&M document 2- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for BIAL for third control period in the MDA_O&M document. 3- Justification provided by BIAL for the repairs and maintenance expense for third control period and the variance of repairs and maintenance expense with authority approved figures in third control period in the  MDA_O&M document. 4- Year on Year growth of actual repairs and maintenance expenditure submitted by DIAL for true up of third control period in the fourth control period consultation paper or tariff order. 5- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for DIAL for third control period in the fourth control period consultation paper or tariff order. 6- Justification for actual repairs and maintenance expense in third control period provided by DIAL and the variance with authority approved figures for the third control period in fourth control period consultation paper or tariff order. 7- Examination and rationale provided by authority on actual repairs and maintenance cost submitted by DIAL for third control period in the fourth control period consultation paper or tariff order. 8- Year on Year growth of actual repairs and maintenance expenditure submitted by MIAL for true up of third control period in the fourth control period consultation paper or tariff order. 9- Justification for actual repairs and maintenance expense in third control period provided by MIAL and the variance with authority approved figures for the third control period in fourth control period consultation paper or tariff order. 10- Year wise repairs and maintenance expenditure as a percentage of regulated asset base for MIAL for third control period in the fourth control period consultation paper or tariff order. 11- Examination and rationale provided by authority on actual repairs and maintenance cost submitted by MIAL for third control period in the fourth control period consultation paper or tariff order. 12- Using the rationale extracted in steps 5, 6, 8, and 9 suggest how the rationale or justification provided by BIAL in the MDA document for repairs and maintenance expenditure for third control period can be enhanced. For every suggestion made, give specific reason why the suggestion is made using relevant references from DIAL and MIAL tariff orders or consultation papers",
                "Analysis of repairs and maintenance expenditure projection for BIAL for fourth control period": f"The uploaded document proposes the following: '{{document_summary}}'. Use the following steps for analysing the repairs and maintenance expenditure projected by BIAL: 1- Year on Year growth of repairs and maintenance expenditure projections by BIAL for fourth control period in the MDA_O&M document 2- Year wise repairs and maintenance expenditure projection as a percentage of regulated asset base for BIAL for fourth control period in the MDA_O&M document. 3- Justification provided by BIAL for the repairs and maintenance expense for fourth control period in the MDA_O&M document. 4- Year on Year growth of repairs and maintenance expenditure projections submitted by DIAL for fourth control period in the fourth control period consultation paper or tariff order. 5- Year wise repairs and maintenance expenditure projections as a percentage of regulated asset base for DIAL for fourth control period in the fourth control period consultation paper or tariff order. 6- Justification for repairs and maintenance expense projections in fourth control period provided by DIAL in fourth control period consultation paper or tariff order. 7- Examination and rationale provided by authority on repairs and maintenance expenditure projections submitted by DIAL for fourth control period in the fourth control period consultation paper or tariff order. 8- Year on Year growth of repairs and maintenance expenditure projections submitted by MIAL for fourth control period in the fourth control period consultation paper or tariff order. 9- Year wise repairs and maintenance expenditure projections as a percentage of regulated asset base for MIAL for fourth control period in the fourth control period consultation paper or tariff order. 10- Justification for repairs and maintenance expense projections in fourth control period provided by MIAL in fourth control period consultation paper or tariff order. 11- Examination and rationale provided by authority on repairs and maintenance expenditure projections submitted by MIAL for fourth control period in the fourth control period consultation paper or tariff order 12- Using the rationale extracted in steps 5, 6, 8, and 9 suggest how the rationale or justification provided by BIAL in the MDA document for repairs and maintenance expenditure for fourth control period can be enhanced. For every suggestion made, give specific reason why the suggestion is made using relevant references from DIAL and MIAL tariff orders or consultation papers",
            },
        }

        st.session_state.mda_analysis_type = st.selectbox(
            "Select Analysis Category:", list(analysis_prompts_config.keys())
        )
        specific_analysis_options = analysis_prompts_config.get(
            st.session_state.mda_analysis_type, {}
        )

        selected_specific_analysis_title = None
        if specific_analysis_options:
            selected_specific_analysis_title = st.selectbox(
                "Select Specific Analysis:", list(specific_analysis_options.keys())
            )

        uploaded_word_file = st.file_uploader(
            "Upload a Word Document (.docx)", type="docx"
        )

        if st.button(
            "Generate Specific Analysis Report",
            type="primary",
            disabled=(not uploaded_word_file or not selected_specific_analysis_title),
        ):
            st.session_state.mda_report_content, st.session_state.mda_chat_history = (
                None,
                [],
            )
            with st.spinner("Processing document and generating analysis..."):
                if (
                    extracted_text := extract_text_from_docx(uploaded_word_file)
                ) and selected_specific_analysis_title:
                    report_parts = [
                        f"## Analysis Report for: *{html.escape(uploaded_word_file.name)}*"
                    ]

                    st.write("Step 1/2: Generating context from uploaded document...")
                    summary_prompt = f"Please provide a detail, neutral summary of the key points of the following document text:\n\n---\n{extracted_text[:20000]}\n---"
                    document_summary = (
                        synthesis_openai_client.chat.completions.create(
                            model=DEPLOYMENT_ID_VAL,
                            messages=[{"role": "user", "content": summary_prompt}],
                        )
                        .choices[0]
                        .message.content
                    )
                    report_parts.append(
                        f"### 1. Document Context (Generated Summary)\n{document_summary}"
                    )

                    st.write(
                        f"Step 2/2: Running analysis for '{selected_specific_analysis_title}'..."
                    )
                    prompt_template = specific_analysis_options.get(
                        selected_specific_analysis_title
                    )

                    if prompt_template:
                        full_prompt = prompt_template.format(
                            document_summary=document_summary
                        )
                        analysis_answer = generate_answer_from_search(
                            user_question=full_prompt,
                            index_name=st.session_state.conv_agent_selected_index,
                            use_hybrid_semantic=st.session_state.conv_agent_use_hybrid,
                            vector_field=st.session_state.conv_agent_vector_field,
                            semantic_config=st.session_state.conv_agent_semantic_config,
                            temperature=st.session_state.conv_agent_temp,
                            max_tokens_param=st.session_state.conv_agent_max_tokens,
                            client_for_synthesis=synthesis_openai_client,
                            show_details=False,
                        )
                        report_parts.append(
                            f"### {selected_specific_analysis_title}\n{analysis_answer}"
                        )

                    st.session_state.mda_report_content = "\n\n---\n\n".join(
                        report_parts
                    )
                    st.success("Analysis report generated successfully!")

        if st.session_state.mda_report_content:
            st.markdown("---")

            word_file = create_word_document(st.session_state.mda_report_content)
            st.download_button(
                label="📥 Download Report as Word",
                data=word_file,
                file_name=f"BIAL_Analysis_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            st.markdown(st.session_state.mda_report_content, unsafe_allow_html=True)

            st.markdown("---")
            st.subheader("💬 Follow-up Chat")
            for i, message in enumerate(st.session_state.mda_chat_history):
                with st.chat_message(message["role"]):
                    st.markdown(message["content"], unsafe_allow_html=True)
                    if message["role"] == "assistant":
                        if st.button("✨ Integrate & Refine Report", key=f"refine_{i}"):
                            st.session_state.mda_report_content = (
                                refine_and_regenerate_report(
                                    st.session_state.mda_report_content,
                                    message["content"],
                                    synthesis_openai_client,
                                )
                            )
                            st.session_state.mda_chat_history = []
                            st.rerun()

            if prompt := st.chat_input("Ask to elaborate or find new information..."):
                st.session_state.mda_chat_history.append(
                    {"role": "user", "content": prompt}
                )
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.spinner("Agent is thinking..."):
                    answer = None
                    prompt_lower = prompt.lower()

                    web_search_keywords = [
                        "bing search",
                        "search bing",
                        "google search",
                        "serper",
                        "search google",
                        "latest",
                        "current",
                        "web search",
                        "internet",
                    ]
                    is_web_search = any(w in prompt_lower for w in web_search_keywords)

                    if is_web_search:
                        st.info(
                            f"⚡ Performing Web Search ({st.session_state.web_search_engine})..."
                        )

                        search_tool = query_bing_web_search
                        if "Google" in st.session_state.web_search_engine:
                            search_tool = query_serpapi

                        context_from_search = search_tool(prompt)

                        # <<< ROBUST ERROR HANDLING >>>
                        if context_from_search.startswith("Error:"):
                            answer = context_from_search  # Directly show the error to the user
                        else:
                            web_search_prompt = f"""You are an expert AI assistant. Based on the following SEARCH RESULTS, provide a concise and informative answer to the user's question.
                            USER QUESTION: "{prompt}"
                            SEARCH RESULTS:\n---\n{context_from_search}\n---\nYOUR DETAILED, FORMATTED ANSWER:"""

                            response = (
                                synthesis_openai_client.chat.completions.create(
                                    model=DEPLOYMENT_ID_VAL,
                                    messages=[
                                        {"role": "user", "content": web_search_prompt}
                                    ],
                                )
                                .choices[0]
                                .message.content
                            )
                            answer = response
                    else:
                        st.info(f"🧠 Searching Internal Documents...")
                        answer = generate_answer_from_search(
                            user_question=prompt,
                            index_name=st.session_state.conv_agent_selected_index,
                            use_hybrid_semantic=st.session_state.conv_agent_use_hybrid,
                            vector_field=st.session_state.conv_agent_vector_field,
                            semantic_config=st.session_state.conv_agent_semantic_config,
                            temperature=st.session_state.conv_agent_temp,
                            max_tokens_param=st.session_state.conv_agent_max_tokens,
                            client_for_synthesis=synthesis_openai_client,
                            show_details=False,
                        )

                    if answer:
                        st.session_state.mda_chat_history.append(
                            {"role": "assistant", "content": answer}
                        )
                        st.rerun()


if __name__ == "__main__":
    try:
        main_app_logic()
    except Exception as e:
        st.error(f"An critical unexpected error occurred: {e}")
        traceback.print_exc()
